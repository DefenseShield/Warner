from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Function to add heading
def add_heading(text, level=0):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(14 if level == 0 else 12)
        run.bold = True

# Function to add paragraph
def add_paragraph(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Function to add table
def add_table(data, headers):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # Headers
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    # Data
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.cell(row_idx, col_idx).text = str(cell_data)

# --- Cover Page ---
add_heading("Planning Report: Military Route Bogotá - Ciudad Juárez", 0)
add_paragraph("\n", size=12)
add_paragraph(f"Date: {datetime.now().strftime('%B %d, 2025')}", bold=True, size=12)
add_paragraph("Purpose: Planning and execution of a military land route from Bogotá, Colombia, to Ciudad Juárez, Mexico, prioritizing speed, safety, and logistics.", size=12)
doc.add_page_break()

# --- Introduction ---
add_heading("Introduction", 1)
add_paragraph(
    "This report outlines the planning of a military land route from Bogotá, Colombia, to Ciudad Juárez, Mexico, covering approximately 4,800 km across multiple countries (Colombia, Panama, Costa Rica, Nicaragua, Honduras, Guatemala, and Mexico). The route is designed for a military convoy, prioritizing main highways, safety against conflict zones, and efficient logistics. Due to the Darién Gap, a maritime or air segment is included between Colombia and Panama. The objective is to ensure a swift and secure execution, with an estimated duration of 7 to 10 days, accounting for stops, border crossings, and resupply."
)

# --- Route Description ---
add_heading("Route Description", 1)
add_paragraph(
    "The route is divided into key segments, utilizing major highways such as the Pan-American Highway and toll roads in Mexico. Below is a detailed breakdown of each segment:"
)

# Route data
route_data = [
    ["1", "Bogotá → Cúcuta, Colombia", "560 km", "Troncal 55, Troncal 66", "10-12 hours", "Border area with Venezuela, potential instability. Military escort recommended."],
    ["2", "Cúcuta → Cartagena, Colombia", "430 km", "Troncal 45, Route 90", "8 hours", "Preparation for maritime/air transport to Panama."],
    ["3", "Cartagena → Panama City", "600 km (gap)", "Maritime/Air", "12-24 hours", "Coordinate with military ports or airports."],
    ["4", "Panama City → San José, Costa Rica", "530 km", "Pan-American Highway (CA-1)", "8-10 hours", "Border crossing at Paso Canoas, strict documentation."],
    ["5", "San José → Managua, Nicaragua", "430 km", "Pan-American Highway (CA-1)", "7-8 hours", "Political risks in Nicaragua, maintain low profile."],
    ["6", "Managua → Tegucigalpa, Honduras", "400 km", "CA-1, CA-3", "7-8 hours", "Areas with criminal activity, use armed escorts."],
    ["7", "Tegucigalpa → Guatemala City", "360 km", "CA-5, CA-13, CA-9", "6-7 hours", "Avoid gang areas, coordinate with Guatemalan army."],
    ["8", "Guatemala City → Tapachula, Mexico", "250 km", "CA-1, Mexico 200", "5-6 hours", "Border with high migratory activity, reinforce security."],
    ["9", "Tapachula → Mexico City", "1,100 km", "Mexico 200, Mexico 150D", "14-16 hours", "Use toll highways, avoid cartel zones."],
    ["10", "Mexico City → Ciudad Juárez", "1,800 km", "Mexico 57D, Mexico 45D", "20-22 hours", "Escorted convoys in Chihuahua, military presence in Juárez."]
]

add_table(route_data, ["No.", "Segment", "Distance", "Highway", "Estimated Time", "Notes"])

# --- Military Considerations ---
add_heading("Military Considerations", 1)

add_heading("Security", 2)
add_paragraph(
    "- Avoid high-risk areas: Darién Gap (impassable by land), urban areas in Honduras (San Pedro Sula), and Mexican regions with cartel presence (Guerrero, Chihuahua outside highways).\n"
    "- Deploy armed convoys with armored vehicles and escorts in vulnerable segments.\n"
    "- Use surveillance drones and satellite communication for real-time monitoring.\n"
    "- Establish protocols for border crossings, with personnel trained in negotiation and crowd control."
)

add_heading("Logistics", 2)
add_paragraph(
    "- Plan resupply stations every 300-400 km, especially on main highways.\n"
    "- Coordinate with local armies for expedited border permits (Paso Canoas, Peñas Blancas, Ciudad Hidalgo).\n"
    "- Secure maritime or air transport from Cartagena to Panama, with capacity for vehicles and personnel.\n"
    "- Maintain reserves of fuel, food, and spare parts in each convoy."
)

# --- Timeline ---
add_heading("Timeline", 1)
add_paragraph(
    "The total estimated time for the route is 7 to 10 days, considering driving, rest, border crossings, and logistics. Below is the approximate timeline:"
)

timeline_data = [
    ["Day 1", "Bogotá → Cúcuta", "560 km", "10-12 hours", "Overnight in Cúcuta, security review."],
    ["Day 2", "Cúcuta → Cartagena", "430 km", "8 hours", "Preparation for Panama transfer."],
    ["Day 3", "Cartagena → Panama City", "600 km (gap)", "12-24 hours", "Maritime/air logistics."],
    ["Day 4", "Panama City → San José", "530 km", "8-10 hours", "Rest in San José."],
    ["Day 5", "San José → Managua", "430 km", "7-8 hours", "Border review in Nicaragua."],
    ["Day 6", "Managua → Tegucigalpa", "400 km", "7-8 hours", "Overnight in Tegucigalpa, enhanced security."],
    ["Day 7", "Tegucigalpa → Guatemala City", "360 km", "6-7 hours", "Coordination with Guatemala."],
    ["Day 8", "Guatemala City → Tapachula", "250 km", "5-6 hours", "Entry to Mexico, customs inspection."],
    ["Day 9", "Tapachula → Mexico City", "1,100 km", "14-16 hours", "Rest in Mexico City."],
    ["Day 10", "Mexico City → Ciudad Juárez", "1,800 km", "20-22 hours", "Arrival at destination."]
]

add_table(timeline_data, ["Day", "Segment", "Distance", "Time", "Notes"])

add_paragraph(
    f"Total driving time: ~95-110 hours (~4-5 days without stops). With rest and logistics, estimated at 7-10 days."
)

# --- Recommendations ---
add_heading("Recommendations", 1)
add_paragraph(
    "- Conduct prior route reconnaissance using satellite intelligence and drones to identify threats.\n"
    "- Establish mobile checkpoints in each country, with constant communication between convoys.\n"
    "- Train personnel in border procedures and ambush response.\n"
    "- Prioritize toll highways in Mexico (150D, 57D, 45D) to minimize risks.\n"
    "- Maintain flexibility to adjust the route in case of road closures or unexpected conflicts."
)

# --- Conclusion ---
add_heading("Conclusion", 1)
add_paragraph(
    "The military route from Bogotá to Ciudad Juárez is feasible in 7 to 10 days, using main highways and a logistical gap in the Darién. Detailed planning, with emphasis on security and logistics, ensures efficient execution. International coordination and the use of surveillance technology will be critical to the operation's success."
)

# Save document
output_path = "Military_Route_Report_Bogota_Juarez.docx"
doc.save(output_path)
print(f"Report saved as '{output_path}'")